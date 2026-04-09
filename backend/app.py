import os
import uuid
import threading
import json
from pathlib import Path
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv
import whisper
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sse_starlette.sse import EventSourceResponse

ENV_PATH = Path(__file__).parent.parent / ".env"
load_dotenv(dotenv_path=str(ENV_PATH))

app = FastAPI(title="Kratos API")

# CORS for frontend dev server
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories
BASE_DIR = Path(__file__).parent
OUTPUTS_DIR = BASE_DIR / "outputs"
TEMP_DIR = BASE_DIR / "temp"
OUTPUTS_DIR.mkdir(exist_ok=True)
TEMP_DIR.mkdir(exist_ok=True)

# Load Whisper model at startup
WHISPER_MODEL_SIZE = os.getenv("WHISPER_MODEL", "small")
print(f"\nLoading Whisper model '{WHISPER_MODEL_SIZE}'...")
print("(This may take a moment on first run while the model downloads.)")
whisper_model = whisper.load_model(WHISPER_MODEL_SIZE)
print(f"Whisper model '{WHISPER_MODEL_SIZE}' loaded successfully.\n")

# Anthropic client — held in a mutable container so /api/settings can reload it
def _is_real_key(key: str) -> bool:
    """Return True only if the key looks like a real value (not a placeholder)."""
    return bool(key) and key != "your_api_key_here" and not key.startswith("your_")

def _make_anthropic_client():
    key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not _is_real_key(key):
        return None
    try:
        return anthropic.Anthropic(api_key=key)
    except Exception:
        return None

anthropic_client = _make_anthropic_client()

def _reload_anthropic_client():
    global anthropic_client
    anthropic_client = _make_anthropic_client()

# Diarization pipeline — loaded lazily on first use so startup isn't slowed
# when diarization is not needed.
_diarization_pipeline = None
_diarization_lock = threading.Lock()


def _get_diarization_pipeline():
    global _diarization_pipeline
    if _diarization_pipeline is not None:
        return _diarization_pipeline
    with _diarization_lock:
        if _diarization_pipeline is not None:
            return _diarization_pipeline
        hf_token = os.getenv("HF_TOKEN", "").strip()
        if not hf_token:
            raise RuntimeError(
                "Hugging Face token is not configured. Add it in the Settings page to use speaker diarization."
            )
        try:
            import torch
            from pyannote.audio import Pipeline
            print("Loading pyannote speaker-diarization pipeline (first use)…")
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=hf_token,
            )
            if pipeline is None or not callable(pipeline):
                raise RuntimeError(
                    "pyannote returned no pipeline. Make sure your HF token is valid and "
                    "you accepted the pyannote/speaker-diarization-3.1 model license at "
                    "https://hf.co/pyannote/speaker-diarization-3.1"
                )
            # Use MPS on Apple Silicon if available, otherwise CPU
            if torch.backends.mps.is_available():
                pipeline.to(torch.device("mps"))
                print("Diarization pipeline loaded (MPS).")
            else:
                pipeline.to(torch.device("cpu"))
                print("Diarization pipeline loaded (CPU).")
            _diarization_pipeline = pipeline
        except Exception as exc:
            raise RuntimeError(f"Failed to load diarization pipeline: {exc}") from exc
    return _diarization_pipeline


# In-memory storage
uploaded_files: dict[str, Path] = {}
transcription_status: dict[str, dict] = {}


# --- Pydantic models ---

class SummarizeRequest(BaseModel):
    transcript: str
    summary_format: str
    custom_prompt: Optional[str] = None
    original_filename: Optional[str] = None


class ExportRequest(BaseModel):
    transcript: str
    summary: str
    original_filename: str
    summary_format: str


class SettingsRequest(BaseModel):
    anthropic_api_key: Optional[str] = None
    hf_token: Optional[str] = None


# --- Background transcription ---

def _make_whisper_tqdm(file_id: str):
    """Return a tqdm subclass that maps Whisper's decode loop to real progress (5–85%)."""
    import tqdm as tqdm_module

    status = transcription_status

    class _WhisperProgress(tqdm_module.tqdm):
        def update(self, n=1):
            super().update(n)
            if self.total:
                pct = 5 + int(self.n / self.total * 80)
                if file_id in status:
                    status[file_id]["progress"] = min(pct, 85)

    return _WhisperProgress


def _align_diarization(whisper_segments: list, diarization) -> str:
    """
    Merge pyannote diarization turns with Whisper word-level segments.

    Strategy: for each Whisper segment find which speaker's turn has the
    greatest overlap with that segment's time range, then group consecutive
    segments that belong to the same speaker into a single labelled block.
    """
    # Build a flat list of (start, end, speaker) turns from pyannote
    turns = [
        (turn.start, turn.end, speaker)
        for turn, _, speaker in diarization.itertracks(yield_label=True)
    ]

    # Sort speakers so labels are stable (SPEAKER_00 → Speaker 1, etc.)
    speaker_order: dict[str, int] = {}

    def _label(spk: str) -> str:
        if spk not in speaker_order:
            speaker_order[spk] = len(speaker_order) + 1
        return f"Speaker {speaker_order[spk]}"

    def _best_speaker(seg_start: float, seg_end: float) -> str:
        best_spk, best_overlap = "Speaker 1", 0.0
        for t_start, t_end, spk in turns:
            overlap = max(0.0, min(seg_end, t_end) - max(seg_start, t_start))
            if overlap > best_overlap:
                best_overlap = overlap
                best_spk = spk
        return _label(best_spk)

    # Assign a speaker to each Whisper segment
    labelled = [
        (_best_speaker(seg["start"], seg["end"]), seg["text"].strip())
        for seg in whisper_segments
        if seg.get("text", "").strip()
    ]

    # Merge consecutive segments with the same speaker
    lines: list[str] = []
    current_speaker, current_parts = None, []
    for speaker, text in labelled:
        if speaker != current_speaker:
            if current_speaker is not None:
                lines.append(f"{current_speaker}: {' '.join(current_parts)}")
            current_speaker = speaker
            current_parts = [text]
        else:
            current_parts.append(text)
    if current_speaker:
        lines.append(f"{current_speaker}: {' '.join(current_parts)}")

    return "\n\n".join(lines)


def _run_transcription(file_id: str, file_path: Path, diarize: bool = False):
    """Run Whisper transcription (and optional diarization) in a background thread."""
    transcription_status[file_id] = {"status": "processing", "progress": 5, "transcript": None, "error": None}

    try:
        import tqdm as tqdm_module
        original_tqdm = tqdm_module.tqdm
        tqdm_module.tqdm = _make_whisper_tqdm(file_id)
        try:
            # Always request word-level segments so diarization alignment can use timestamps
            result = whisper_model.transcribe(str(file_path), verbose=False)
        finally:
            tqdm_module.tqdm = original_tqdm

        if diarize:
            transcription_status[file_id]["progress"] = 90
            pipeline = _get_diarization_pipeline()
            if pipeline is None:
                raise RuntimeError(
                    "Speaker diarization is unavailable. Check your Hugging Face token and "
                    "confirm you accepted the pyannote model license."
                )
            # pyannote only reads WAV — convert if needed
            audio_path = file_path
            tmp_wav = None
            if file_path.suffix.lower() != ".wav":
                import subprocess, tempfile
                tmp_wav = Path(tempfile.mktemp(suffix=".wav"))
                subprocess.run(
                    ["ffmpeg", "-y", "-i", str(file_path), "-ar", "16000", "-ac", "1", str(tmp_wav)],
                    check=True,
                    capture_output=True,
                )
                audio_path = tmp_wav
            try:
                diarization = pipeline(str(audio_path))
            except TypeError as exc:
                raise RuntimeError(
                    "Diarization pipeline failed to run — this usually means the model weights "
                    "were not fully downloaded. Visit https://hf.co/pyannote/speaker-diarization-3.1 "
                    "to accept the license, then restart the server."
                ) from exc
            except Exception as exc:
                # MPS can run out of memory on long files — retry on CPU
                if "mps" in str(exc).lower() or "metal" in str(exc).lower() or "out of memory" in str(exc).lower():
                    print(f"MPS diarization failed ({exc}), retrying on CPU…")
                    import torch
                    pipeline.to(torch.device("cpu"))
                    diarization = pipeline(str(audio_path))
                else:
                    raise
            finally:
                if tmp_wav and tmp_wav.exists():
                    tmp_wav.unlink()
            transcript_text = _align_diarization(result.get("segments", []), diarization)
        else:
            transcript_text = result["text"]

        transcription_status[file_id].update(
            {
                "status": "complete",
                "progress": 100,
                "transcript": transcript_text,
            }
        )
    except Exception as exc:
        transcription_status[file_id].update({"status": "error", "error": str(exc)})
    finally:
        try:
            file_path.unlink(missing_ok=True)
        except Exception:
            pass


# --- Env helpers ---

def _write_env(updates: dict) -> None:
    """Update or insert key=value pairs in the project .env file."""
    lines: list[str] = []
    if ENV_PATH.exists():
        lines = ENV_PATH.read_text(encoding="utf-8").splitlines()

    updated: set[str] = set()
    new_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith("#") and "=" in stripped:
            key = stripped.split("=", 1)[0].strip()
            if key in updates:
                new_lines.append(f"{key}={updates[key]}")
                updated.add(key)
                continue
        new_lines.append(line)

    for key, value in updates.items():
        if key not in updated:
            new_lines.append(f"{key}={value}")

    ENV_PATH.write_text("\n".join(new_lines) + "\n", encoding="utf-8")


# --- Endpoints ---

def _key_status() -> dict:
    """Read current key presence directly from env vars (populated via load_dotenv)."""
    return {
        "anthropic_key_set": bool(os.getenv("ANTHROPIC_API_KEY", "").strip()),
        "hf_token_set": bool(os.getenv("HF_TOKEN", "").strip()),
    }


@app.get("/api/settings")
async def get_settings():
    """Return whether each secret key is currently configured (never the values)."""
    return _key_status()


@app.post("/api/settings")
async def save_settings(req: SettingsRequest):
    """Write non-empty key values to .env and reload them without a server restart."""
    updates: dict[str, str] = {}
    if req.anthropic_api_key is not None:
        updates["ANTHROPIC_API_KEY"] = req.anthropic_api_key.strip()
    if req.hf_token is not None:
        updates["HF_TOKEN"] = req.hf_token.strip()

    if not updates:
        raise HTTPException(status_code=400, detail="No values provided.")

    # Create .env from example if it doesn't exist yet
    if not ENV_PATH.exists():
        example = ENV_PATH.parent / ".env.example"
        if example.exists():
            import shutil
            shutil.copy(str(example), str(ENV_PATH))
        else:
            ENV_PATH.write_text("", encoding="utf-8")

    try:
        _write_env(updates)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to write .env: {exc}")

    # Reload env vars in-process so the server picks them up immediately
    load_dotenv(dotenv_path=str(ENV_PATH), override=True)
    _reload_anthropic_client()

    return _key_status()


@app.post("/api/upload")
async def upload_audio(file: UploadFile = File(...)):
    allowed_exts = {".wav", ".mp3", ".m4a", ".webm", ".ogg", ".flac", ".mp4"}
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in allowed_exts:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{suffix}'. Allowed: {', '.join(allowed_exts)}")

    file_id = str(uuid.uuid4())
    file_path = TEMP_DIR / f"{file_id}{suffix}"

    content = await file.read()
    file_path.write_bytes(content)

    uploaded_files[file_id] = file_path
    transcription_status[file_id] = {"status": "uploaded", "progress": 0, "transcript": None, "error": None}

    return {"file_id": file_id, "filename": file.filename, "size": len(content)}


@app.get("/api/transcribe/{file_id}")
async def transcribe_stream(file_id: str, diarize: bool = Query(False)):
    """Start transcription (if needed) and stream progress via Server-Sent Events."""
    if file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="File not found. Please upload the audio first.")

    current = transcription_status.get(file_id, {})
    if current.get("status") not in ("processing", "complete", "error"):
        file_path = uploaded_files[file_id]
        thread = threading.Thread(
            target=_run_transcription,
            args=(file_id, file_path, diarize),
            daemon=True,
        )
        thread.start()

    async def event_generator():
        import asyncio
        while True:
            status = transcription_status.get(file_id, {})
            yield {"event": "status", "data": json.dumps(status)}

            if status.get("status") in ("complete", "error"):
                break

            await asyncio.sleep(0.5)

    return EventSourceResponse(event_generator())


@app.post("/api/summarize")
async def summarize_transcript(request: SummarizeRequest):
    if not anthropic_client:
        raise HTTPException(
            status_code=400,
            detail="Anthropic API key is not configured. Add it in the Settings page.",
        )

    format_prompts = {
        "meeting_notes": (
            "You are a professional meeting notes writer. "
            "Analyze the following transcript and produce structured meeting notes with these sections:\n"
            "## Attendees\n(List any names or roles mentioned; write 'Not specified' if unclear)\n\n"
            "## Key Discussion Points\n(Bullet points of main topics covered)\n\n"
            "## Decisions Made\n(Any conclusions or agreements reached)\n\n"
            "## Action Items\n(Tasks with owners and deadlines if mentioned)\n\n"
            "## Follow-ups\n(Open questions or next steps)\n\n"
            "Be concise but thorough. Use clear, professional language."
        ),
        "key_takeaways": (
            "You are a professional summarizer. "
            "Analyze the following transcript and produce two sections:\n\n"
            "## Key Takeaways\n"
            "List the most important insights as bullet points. Each bullet must follow this exact format:\n"
            "• **Title:** Description\n"
            "Where Title is a short bold label (2-5 words) and Description is one clear sentence explaining the point.\n\n"
            "## Action Items\n"
            "List every action item, task, or commitment mentioned as a markdown table with exactly these columns:\n"
            "| Item | Owner | Timeline |\n"
            "| --- | --- | --- |\n"
            "Fill in Owner and Timeline from the transcript; write 'TBD' if not mentioned. "
            "If there are no action items, write 'No action items identified.'"
        ),
        "executive_summary": (
            "You are a professional executive assistant. "
            "Write a polished 2-3 paragraph executive summary of the following transcript. "
            "This summary should be suitable for forwarding to senior leadership. "
            "Focus on strategic implications, key outcomes, and important context. "
            "Use clear, professional prose with no bullet points."
        ),
        "follow_up_email": (
            "You are a professional business writer. "
            "Draft a concise follow-up email based on the following transcript. "
            "The email should:\n"
            "- Have a clear subject line on the first line prefixed with 'Subject:'\n"
            "- Open with a brief thank-you or context line\n"
            "- Summarize the key points discussed in 2-4 short paragraphs\n"
            "- List any action items or next steps clearly\n"
            "- Close professionally\n"
            "Use clear, professional language suitable for sending to all participants."
        ),
        "custom": request.custom_prompt or "Summarize the following transcript clearly and concisely.",
    }

    system_prompt = format_prompts.get(request.summary_format, format_prompts["key_takeaways"])
    # If the user entered custom instructions, append them to the chosen format's prompt
    if request.custom_prompt and request.summary_format != "custom":
        system_prompt = system_prompt + "\n\nAdditional instructions from the user: " + request.custom_prompt

    try:
        message = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": f"Please summarize the following transcript:\n\n{request.transcript}",
                }
            ],
        )
        summary_text = message.content[0].text

        # Auto-save transcript as Word doc
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        stem = Path(request.original_filename or "recording").stem
        safe_stem = "".join(c for c in stem if c.isalnum() or c in (" ", "-", "_")).strip() or "recording"
        t_doc = Document()
        _apply_doc_style(t_doc)
        h = t_doc.add_heading("Transcript", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_meta_line(t_doc, f"Generated: {timestamp}  |  Source: {request.original_filename or 'recording'}")
        t_doc.add_paragraph()
        for line in request.transcript.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            p = t_doc.add_paragraph()
            if ":" in stripped:
                prefix, _, rest = stripped.partition(":")
                if prefix.strip().startswith("Speaker ") and prefix.strip()[8:].strip().isdigit():
                    label_run = p.add_run(prefix.strip() + ":")
                    label_run.bold = True
                    label_run.font.name = "Calibri"
                    label_run.font.size = Pt(11)
                    text_run = p.add_run(" " + rest.strip())
                    text_run.font.name = "Calibri"
                    text_run.font.size = Pt(11)
                    continue
            run = p.add_run(stripped)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        transcript_filename = f"{safe_stem}_transcript.docx"
        t_doc.save(str(OUTPUTS_DIR / transcript_filename))

        return {"summary": summary_text}
    except anthropic.AuthenticationError:
        raise HTTPException(status_code=401, detail="Invalid ANTHROPIC_API_KEY. Please check your .env file.")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Summarization failed: {exc}")


@app.post("/api/export")
async def export_documents(request: ExportRequest):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    stem = Path(request.original_filename).stem
    safe_stem = "".join(c for c in stem if c.isalnum() or c in (" ", "-", "_")).strip() or "recording"

    format_labels = {
        "meeting_notes": "Meeting Notes",
        "key_takeaways": "Key Takeaways",
        "executive_summary": "Executive Summary",
        "follow_up_email": "Follow-up Email",
        "custom": "Custom Summary",
    }
    format_label = format_labels.get(request.summary_format, "Summary")

    # --- Transcript document ---
    t_doc = Document()
    _apply_doc_style(t_doc)

    h = t_doc.add_heading("Transcript", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_meta_line(t_doc, f"Generated: {timestamp}  |  Source: {request.original_filename}")
    t_doc.add_paragraph()

    for line in request.transcript.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        p = t_doc.add_paragraph()
        # Bold the "Speaker N:" prefix if present
        if ":" in stripped:
            prefix, _, rest = stripped.partition(":")
            if prefix.strip().startswith("Speaker ") and prefix.strip()[8:].strip().isdigit():
                label_run = p.add_run(prefix.strip() + ":")
                label_run.bold = True
                label_run.font.name = "Calibri"
                label_run.font.size = Pt(11)
                text_run = p.add_run(" " + rest.strip())
                text_run.font.name = "Calibri"
                text_run.font.size = Pt(11)
                continue
        run = p.add_run(stripped)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    transcript_filename = f"{safe_stem}_transcript.docx"
    t_doc.save(str(OUTPUTS_DIR / transcript_filename))

    # --- Summary document ---
    s_doc = Document()
    _apply_doc_style(s_doc)

    h = s_doc.add_heading(format_label, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_meta_line(
        s_doc,
        f"Generated: {timestamp}  |  Format: {format_label}  |  Source: {request.original_filename}",
    )
    s_doc.add_paragraph()

    for line in request.summary.split("\n"):
        stripped = line.rstrip()
        if not stripped:
            s_doc.add_paragraph()
            continue
        if stripped.startswith("## "):
            s_doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            s_doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("### "):
            s_doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = s_doc.add_paragraph(stripped[2:], style="List Bullet")
            _set_normal_font(p)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = s_doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(12)
        else:
            p = s_doc.add_paragraph(stripped)
            _set_normal_font(p)

    summary_filename = f"{safe_stem}_summary.docx"
    s_doc.save(str(OUTPUTS_DIR / summary_filename))

    return {
        "transcript_url": f"/api/download/{transcript_filename}",
        "summary_url": f"/api/download/{summary_filename}",
        "transcript_filename": transcript_filename,
        "summary_filename": summary_filename,
    }


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    safe_name = Path(filename).name  # prevent path traversal
    file_path = OUTPUTS_DIR / safe_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")

    return FileResponse(
        path=str(file_path),
        filename=safe_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# --- Document helpers ---

def _apply_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _add_meta_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _set_normal_font(para):
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


# --- Serve built frontend in production ---
_frontend_dist = BASE_DIR.parent / "frontend" / "dist"
if _frontend_dist.exists():
    app.mount("/", StaticFiles(directory=str(_frontend_dist), html=True), name="static")


if __name__ == "__main__":
    import uvicorn
    host = os.getenv("HOST", "127.0.0.1")
    port = int(os.getenv("BACKEND_PORT", "8000"))
    print(f"Starting Kratos backend on http://{host}:{port}")
    uvicorn.run("app:app", host=host, port=port, reload=False)
